from docx import Document

# Créer un document Word
doc = Document()

# Titre du document
doc.add_heading('Rapport de Projet – Application de Gestion des Projets', 0)

# 1. Informations générales
doc.add_heading('1. Informations générales', level=1)
doc.add_paragraph('Nom du projet : Gestion des Projets')
doc.add_paragraph('Technologies utilisées : Python, Flask, MongoDB, Bootstrap 5')
doc.add_paragraph('Date de création : Janvier 2026')
doc.add_paragraph('Auteur / Équipe : [Nom / Équipe]')

# 2. Objectifs
doc.add_heading('2. Objectifs du projet', level=1)
doc.add_paragraph(
    'L’objectif principal est de créer une application web permettant :\n'
    '- La gestion centralisée de projets et de leurs tâches.\n'
    '- L’organisation efficace des tâches par état, priorité et échéance.\n'
    '- Une interface moderne et intuitive pour l’utilisateur.\n'
    '- La possibilité d’extension future avec des comptes utilisateurs, filtres, pagination et drag & drop.'
)

# 3. Architecture
doc.add_heading('3. Architecture de l’application', level=1)
doc.add_heading('3.1 Structure du projet', level=2)
doc.add_paragraph(
    '/gestion_projets\n'
    '│\n'
    '├─ app.py\n'
    '├─ requirements.txt\n'
    '├─ /templates\n'
    '│   ├─ base.html\n'
    '│   ├─ index.html\n'
    '│   ├─ add_project.html\n'
    '│   ├─ edit_project.html\n'
    '│   ├─ project_detail.html\n'
    '├─ /static\n'
    '│   ├─ /css\n'
    '│   │   └─ style.css\n'
    '│   └─ /js\n'
    '│       └─ script.js'
)
doc.add_heading('3.2 Schéma MongoDB', level=2)
doc.add_paragraph(
    'Collection `projets` :\n'
    '{\n'
    '    "_id": ObjectId,\n'
    '    "nom": "Nom du projet",\n'
    '    "description": "Description du projet",\n'
    '    "taches": [\n'
    '        {\n'
    '            "nom": "Nom de la tâche",\n'
    '            "etat": "À faire | En cours | Terminé",\n'
    '            "priorite": "Haute | Moyenne | Basse",\n'
    '            "echeance": "YYYY-MM-DD"\n'
    '        }\n'
    '    ]\n'
    '}'
)

# 4. Fonctionnalités principales
doc.add_heading('4. Fonctionnalités principales', level=1)
doc.add_paragraph(
    '1. Gestion des projets\n'
    '- Création, modification et suppression de projets.\n'
    '- Liste des projets avec actions rapides (voir, modifier, supprimer).\n\n'
    '2. Gestion des tâches par projet\n'
    '- Ajout, modification et suppression de tâches.\n'
    '- Chaque tâche possède : Nom, État, Priorité, Échéance.\n'
    '- Modification des tâches via un formulaire modal.\n'
    '- Affichage dynamique avec Bootstrap et couleurs indicatives selon priorité.\n\n'
    '3. Interface moderne\n'
    '- Thème sombre uniforme (dark mode).\n'
    '- Formulaires modernes et responsives avec Bootstrap 5.\n'
    '- Boutons avec icônes et labels clairs.'
)

# 5. Expérience utilisateur
doc.add_heading('5. Expérience utilisateur', level=1)
doc.add_paragraph(
    '- Tableau clair pour visualiser rapidement les projets et tâches.\n'
    '- Formulaires intuitifs avec labels, placeholders et validation HTML5.\n'
    '- Modale pour modification des tâches, améliorant la fluidité et réduisant les clics.\n'
    '- Flash messages pour les retours d’action (succès, erreur).'
)

# 6. Extensions prévues
doc.add_heading('6. Extensions prévues', level=1)
doc.add_paragraph(
    '- Comptes utilisateurs : chaque utilisateur gère ses projets de manière sécurisée.\n'
    '- Filtrage et recherche de tâches : trouver rapidement une tâche par état, priorité ou échéance.\n'
    '- Pagination des projets : navigation optimisée pour de nombreux projets.\n'
    '- Drag & drop des tâches : réorganisation intuitive.\n'
    '- Notifications / rappels d’échéance : alertes visuelles ou par email.'
)

# 7. Technologies et librairies
doc.add_heading('7. Technologies et librairies', level=1)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Technologie'
hdr_cells[1].text = 'Usage'
rows = [
    ('Python 3', 'Langage principal'),
    ('Flask', 'Framework web'),
    ('Flask-PyMongo', 'Connexion à MongoDB'),
    ('MongoDB', 'Base de données NoSQL'),
    ('Bootstrap 5', 'UI moderne et responsive'),
    ('Flask-Bcrypt', 'Hachage des mots de passe (future extension)'),
    ('Jinja2', 'Templating HTML')
]
for tech, usage in rows:
    row_cells = table.add_row().cells
    row_cells[0].text = tech
    row_cells[1].text = usage

# 8. Diagrammes et captures d’écran
doc.add_heading('8. Diagrammes et captures d’écran', level=1)
doc.add_paragraph(
    '- Diagramme UML des classes / collections.\n'
    '- Exemple d’écran principal avec liste des projets.\n'
    '- Exemple de modal de modification des tâches.\n'
    '(Remplacer par des captures d’écran générées à partir de l’application réelle)'
)

# 9. Conclusion
doc.add_heading('9. Conclusion', level=1)
doc.add_paragraph(
    'Le projet permet déjà une gestion efficace des projets et tâches avec une interface moderne.\n'
    'Il est scalable et extensible, notamment pour : comptes utilisateurs, notifications, filtres et drag & drop.\n'
    'Ce projet constitue une bonne base pour une application collaborative multi-utilisateurs.'
)

# Sauvegarde
doc.save("Rapport_Gestion_Projets.docx")
print("Rapport généré : Rapport_Gestion_Projets.docx")
